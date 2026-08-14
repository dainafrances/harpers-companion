from __future__ import annotations

import csv
import io
import os
from dataclasses import dataclass

from docx import Document
from pypdf import PdfReader


MAX_DOCUMENT_BYTES = int(os.getenv("MAX_DOCUMENT_BYTES", str(8 * 1024 * 1024)))
MAX_DOCUMENT_CHARS = int(os.getenv("MAX_DOCUMENT_CHARS", "24000"))
SUPPORTED_SUFFIXES = {".txt", ".md", ".csv", ".pdf", ".docx"}


@dataclass(frozen=True)
class ExtractedDocument:
    filename: str
    text: str


def is_supported_document(*, filename: str, content_type: str | None) -> bool:
    suffix = os.path.splitext(filename.lower())[1]
    return suffix in SUPPORTED_SUFFIXES or (content_type or "").lower() in {
        "text/plain",
        "text/markdown",
        "text/csv",
        "application/pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }


def extract_document_text(*, filename: str, data: bytes) -> ExtractedDocument:
    if len(data) > MAX_DOCUMENT_BYTES:
        raise ValueError(f"{filename} is larger than the document limit.")

    suffix = os.path.splitext(filename.lower())[1]
    if suffix == ".pdf":
        text = "\n".join(page.extract_text() or "" for page in PdfReader(io.BytesIO(data)).pages)
    elif suffix == ".docx":
        document = Document(io.BytesIO(data))
        text = "\n".join(paragraph.text for paragraph in document.paragraphs)
    elif suffix == ".csv":
        rows = csv.reader(io.StringIO(data.decode("utf-8", errors="replace")))
        text = "\n".join(" | ".join(row) for row in rows)
    elif suffix in {".txt", ".md"}:
        text = data.decode("utf-8", errors="replace")
    else:
        raise ValueError(f"{filename} is not a supported document type.")

    text = text.strip()
    if len(text) > MAX_DOCUMENT_CHARS:
        text = text[:MAX_DOCUMENT_CHARS] + "\n[Document truncated at configured limit.]"
    return ExtractedDocument(filename=filename, text=text or "[No readable text found.]")
